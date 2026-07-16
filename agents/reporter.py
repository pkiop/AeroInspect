"""ReportAgent — 한국어 .docx 형상 점검 보고서 생성.

파이프라인 산출물(:class:`~core.schemas.InspectionItem` 목록)을 받아

1) Gemini 구조화 출력(:class:`~core.schemas.ReportNarrative`)으로
   점검 개요·종합 의견 서술을 생성하고 (이미지 없이 항목 요약 텍스트만 컨텍스트로 사용),
2) python-docx로 표지 / 점검 개요 / 결함 요약 / 항목별 상세 / 종합 의견
   구조의 .docx 보고서를 조립한다.

유형·심각도·조치·좌우 한국어 라벨 매핑은 UI에서 재사용할 수 있도록
모듈 상수로 노출한다.
"""

from __future__ import annotations

import datetime as dt
import io
import logging
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import _Cell

from core import config, imaging, llm
from core.schemas import Discrepancy, InspectionItem, PartRecord, ReportNarrative

logger = logging.getLogger("aeroinspect.reporter")

# ---------------------------------------------------------------------------
# 한국어 라벨 매핑 (보고서/UI 공용 — 모듈 상수로 export)
# ---------------------------------------------------------------------------

#: 결함 유형 → 한국어 라벨
TYPE_LABELS_KO: dict[str, str] = {
    "missing_part": "부품 누락",
    "misinstalled_part": "부품 오장착",
    "foreign_object": "이물질(FOD)",
    "surface_damage": "표면 손상",
    "other": "기타",
}

#: 심각도 → 한국어 라벨
SEVERITY_LABELS_KO: dict[str, str] = {
    "low": "낮음",
    "medium": "중간",
    "high": "높음",
    "critical": "심각",
}

#: 누락 시 조치(disposition) → 한국어 라벨
DISPOSITION_LABELS_KO: dict[str, str] = {
    "ground_aircraft": "비행 금지(운항 중지)",
    "install_before_flight": "비행 전 장착",
    "monitor": "상태 감시(모니터링)",
    "engineering_review": "기술 검토 필요",
}

#: 좌/우 위치 → 한국어 라벨 (aircraft_side / catalog_side 공용)
SIDE_LABELS_KO: dict[str, str] = {
    "left": "좌측",
    "right": "우측",
    "center": "중앙",
    "unknown": "미상",
    "na": "해당 없음",
}

#: 보고서 말미 고정 면책 문구
DISCLAIMER_TEXT: str = (
    "본 보고서는 가상 데이터 기반 데모 산출물로, 실제 감항성 판단에 사용할 수 없습니다."
)

#: 상세 섹션 사진 폭 (2열 나란히 배치 기준)
IMAGE_WIDTH_INCHES: float = 2.8

#: 결함 요약 표 헤더
_SUMMARY_HEADERS: tuple[str, ...] = ("ID", "유형", "부품명", "좌우", "심각도", "조치", "플래그")

_SYSTEM_INSTRUCTION: str = (
    "너는 항공정비 형상 점검 보고서의 서술부(점검 개요, 종합 의견)를 작성하는 "
    "기술 문서 작성자다. 담백하고 사실 중심적인 문체로, 제공된 점검 항목 요약에 "
    "근거해서만 서술하며 추측이나 과장을 하지 않는다. "
    "이 점검은 항공기 축소 모형을 대상으로 한 가상 데이터 기반 데모임을 인지하고 "
    "작성한다. 모든 문장은 격식 있는 한국어 보고서 문체로 쓴다."
)


def _flag_descriptions() -> dict[str, str]:
    """agents.validator의 FLAG_DESCRIPTIONS를 지연 임포트로 가져온다.

    모듈 로드 순서 의존을 피하기 위해 사용 시점에 임포트한다.
    """
    from agents.validator import FLAG_DESCRIPTIONS

    return FLAG_DESCRIPTIONS


# ---------------------------------------------------------------------------
# 서술 생성 컨텍스트
# ---------------------------------------------------------------------------


def _item_summary_line(item: InspectionItem) -> str:
    """서술 생성 컨텍스트용 항목 1건 요약 텍스트(이미지 미포함)."""
    d = item.discrepancy
    record = item.effective_record
    flags = ", ".join(item.validation.flags) if item.validation.flags else "없음"
    return (
        f"- ID={d.discrepancy_id}"
        f" | 유형={TYPE_LABELS_KO.get(d.discrepancy_type, d.discrepancy_type)}"
        f" | 부품명={d.component_name_ko}"
        f" | 좌우={SIDE_LABELS_KO.get(d.aircraft_side, d.aircraft_side)}"
        f" | 심각도={SEVERITY_LABELS_KO.get(d.severity, d.severity)}"
        f" | 조치={DISPOSITION_LABELS_KO.get(record.disposition_if_missing, record.disposition_if_missing)}"
        f" | 플래그={flags}"
        f" | 근거={d.evidence}"
    )


def _build_narrative_context(items: list[InspectionItem], inspector_name: str) -> str:
    """서술(개요/종합 의견) 생성용 텍스트 컨텍스트를 조립한다."""
    lines = [
        "항공기 축소 모형 형상 비교 점검(기준 사진 대비 점검 사진) 결과다.",
        f"점검자: {inspector_name}",
        f"탐지된 구성 차이: {len(items)}건",
    ]
    if items:
        lines.append("항목 요약:")
        lines.extend(_item_summary_line(item) for item in items)
    else:
        lines.append(
            "기준 형상과 점검 형상 간 구성 차이 미발견(이상 없음). "
            "이상이 발견되지 않았다는 취지로 점검 개요와 종합 의견을 작성하라."
        )
    lines.append("위 내용을 바탕으로 점검 개요(overview)와 종합 의견(overall_opinion)을 작성하라.")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# docx 조립 헬퍼
# ---------------------------------------------------------------------------


def _apply_korean_base_font(doc: DocumentObject) -> None:
    """기본(Normal) 스타일에 한글 글꼴을 지정한다 (실패해도 치명적이지 않음)."""
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Malgun Gothic"
        normal.font.size = Pt(10.5)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "Malgun Gothic")
    except Exception as exc:  # noqa: BLE001 — 폰트 설정 실패는 무시 가능
        logger.warning("보고서 기본 글꼴 설정 실패: %s", exc)


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    """표 셀에 텍스트를 채운다 (굵기/정렬 옵션)."""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_cover(doc: DocumentObject, inspector_name: str, inspected_at: dt.datetime) -> None:
    """표지 — 제목과 점검 메타데이터 표."""
    title = doc.add_heading("항공기 형상 점검 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_rows: tuple[tuple[str, str], ...] = (
        ("대상", "항공기 축소 모형(데모)"),
        ("점검일시", inspected_at.strftime("%Y-%m-%d %H:%M")),
        ("점검자", inspector_name),
        ("점검 유형", "형상 비교 점검"),
    )
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    for row, (key, value) in zip(table.rows, meta_rows):
        _set_cell_text(row.cells[0], key, bold=True)
        _set_cell_text(row.cells[1], value)
    doc.add_paragraph()


def _summary_row_values(item: InspectionItem) -> tuple[str, ...]:
    """결함 요약 표 1행 값 (effective_record 기준 조치 표기)."""
    d = item.discrepancy
    record = item.effective_record
    return (
        d.discrepancy_id,
        TYPE_LABELS_KO.get(d.discrepancy_type, d.discrepancy_type),
        d.component_name_ko,
        SIDE_LABELS_KO.get(d.aircraft_side, d.aircraft_side),
        SEVERITY_LABELS_KO.get(d.severity, d.severity),
        DISPOSITION_LABELS_KO.get(record.disposition_if_missing, record.disposition_if_missing),
        ", ".join(item.validation.flags) if item.validation.flags else "-",
    )


def _add_summary_section(doc: DocumentObject, items: list[InspectionItem]) -> None:
    """'2. 결함 요약' — 표 형태 요약 (비어 있으면 '발견된 결함 없음' 한 행)."""
    doc.add_heading("2. 결함 요약", level=1)
    table = doc.add_table(rows=1, cols=len(_SUMMARY_HEADERS))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, _SUMMARY_HEADERS):
        _set_cell_text(cell, header, bold=True, center=True)

    if not items:
        row = table.add_row()
        merged = row.cells[0].merge(row.cells[-1])
        _set_cell_text(merged, "발견된 결함 없음", center=True)
        return

    for item in items:
        row = table.add_row()
        for cell, value in zip(row.cells, _summary_row_values(item)):
            _set_cell_text(cell, value)


def _add_image_cell(cell: _Cell, png_bytes: bytes, caption: str) -> None:
    """상세 표 셀에 어노테이션 사진 + 캡션을 넣는다."""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(IMAGE_WIDTH_INCHES))
    caption_p = cell.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.bold = True


def _add_item_images(
    doc: DocumentObject,
    discrepancy: Discrepancy,
    baseline_image: bytes | None,
    inspection_image: bytes | None,
) -> None:
    """기준/점검 사진을 bbox 오버레이 후 2열 표로 나란히 배치한다."""
    if baseline_image is None or inspection_image is None:
        doc.add_paragraph("첨부 이미지 없음")
        return
    baseline_png = imaging.annotate_baseline(baseline_image, discrepancy.bbox_baseline)
    inspection_png = imaging.annotate_inspection(inspection_image, discrepancy.bbox_inspection)
    table = doc.add_table(rows=1, cols=2)
    _add_image_cell(table.rows[0].cells[0], baseline_png, "기준(정상)")
    _add_image_cell(table.rows[0].cells[1], inspection_png, "점검(현재)")


def _add_part_info(doc: DocumentObject, record: PartRecord) -> None:
    """부품 정보(P/N, 좌/우, Flight Critical 여부) 단락."""
    doc.add_paragraph(
        "부품 정보: "
        f"P/N {record.part_number or '미확인'}"
        f" · 장착 위치 {SIDE_LABELS_KO.get(record.catalog_side, record.catalog_side)}"
        f" · Flight Critical {'예' if record.flight_critical else '아니오'}"
    )


def _add_reference_docs(doc: DocumentObject, record: PartRecord) -> None:
    """근거 인용 목록 — 비어 있으면 '근거 미확정' 굵게 표시."""
    paragraph = doc.add_paragraph()
    paragraph.add_run("근거 인용: ")
    if not record.reference_docs:
        paragraph.add_run("근거 미확정").bold = True
        return
    for ref in record.reference_docs:
        doc.add_paragraph(f"{ref.doc} — {ref.section}", style="List Bullet")


def _add_installation_steps(doc: DocumentObject, record: PartRecord) -> None:
    """재장착 절차 번호 목록 (항목별 1부터 시작하도록 수동 번호 표기)."""
    if not record.installation_steps:
        doc.add_paragraph("재장착 절차: 해당 정보 없음")
        return
    doc.add_paragraph("재장착 절차:")
    for order, step in enumerate(record.installation_steps, start=1):
        step_p = doc.add_paragraph(f"{order}. {step}")
        step_p.paragraph_format.left_indent = Inches(0.25)


def _add_validation_flags(doc: DocumentObject, item: InspectionItem) -> None:
    """검증 플래그 — FLAG_DESCRIPTIONS로 한국어 설명 병기."""
    flags = item.validation.flags
    if not flags:
        doc.add_paragraph("검증 플래그: 없음(규칙 검증 통과)")
        return
    descriptions = _flag_descriptions()
    doc.add_paragraph("검증 플래그:")
    for flag in flags:
        doc.add_paragraph(
            f"{flag} — {descriptions.get(flag, '설명 미등록')}", style="List Bullet"
        )


def _add_item_detail(
    doc: DocumentObject,
    index: int,
    item: InspectionItem,
    baseline_image: bytes | None,
    inspection_image: bytes | None,
) -> None:
    """'3.x [ID] 부품명' 항목별 상세 블록."""
    d = item.discrepancy
    record = item.effective_record
    doc.add_heading(f"3.{index} [{d.discrepancy_id}] {d.component_name_ko}", level=2)
    _add_item_images(doc, d, baseline_image, inspection_image)
    doc.add_paragraph(f"판독 근거: {d.evidence} (이미지상 위치: {d.image_position_desc})")
    _add_part_info(doc, record)
    _add_reference_docs(doc, record)
    _add_installation_steps(doc, record)
    _add_validation_flags(doc, item)


def _add_disclaimer(doc: DocumentObject) -> None:
    """말미 고정 면책 문구."""
    doc.add_paragraph()
    run = doc.add_paragraph().add_run(DISCLAIMER_TEXT)
    run.italic = True


# ---------------------------------------------------------------------------
# ReportAgent
# ---------------------------------------------------------------------------


class ReportAgent:
    """점검 결과를 한국어 .docx 보고서로 산출하는 에이전트.

    Args:
        model: 서술 생성에 사용할 Gemini 모델명 (하드코딩 금지 — 주입).
    """

    def __init__(self, model: str) -> None:
        self._model = model

    def _generate_narrative(
        self, items: list[InspectionItem], inspector_name: str
    ) -> ReportNarrative:
        """항목 요약 텍스트만으로 개요/종합 의견 서술을 생성한다."""
        context = _build_narrative_context(items, inspector_name)
        narrative = llm.generate_structured(
            agent="reporter",
            model=self._model,
            contents=[context],
            response_schema=ReportNarrative,
            system_instruction=_SYSTEM_INSTRUCTION,
            temperature=0.2,
        )
        if not isinstance(narrative, ReportNarrative):
            narrative = ReportNarrative.model_validate(narrative)
        return narrative

    def build_report(
        self,
        items: list[InspectionItem],
        baseline_images: list[bytes],
        inspection_images: list[bytes],
        inspector_name: str,
        output_dir: Path | None = None,
    ) -> tuple[Path, ReportNarrative]:
        """서술 생성 후 .docx 보고서를 조립해 저장한다.

        bbox 오버레이는 각 이미지 리스트의 첫 번째 이미지에 적용한다.
        items가 비어 있어도 '이상 없음' 취지의 보고서를 생성한다
        (요약표 '발견된 결함 없음', 상세 섹션 생략).

        Returns:
            (생성된 .docx 경로, 생성된 ReportNarrative)
        """
        narrative = self._generate_narrative(items, inspector_name)
        now = dt.datetime.now()
        baseline_image = baseline_images[0] if baseline_images else None
        inspection_image = inspection_images[0] if inspection_images else None

        doc = Document()
        _apply_korean_base_font(doc)
        _add_cover(doc, inspector_name, now)

        doc.add_heading("1. 점검 개요", level=1)
        doc.add_paragraph(narrative.overview)

        _add_summary_section(doc, items)

        if items:
            doc.add_heading("3. 항목별 상세", level=1)
            for index, item in enumerate(items, start=1):
                _add_item_detail(doc, index, item, baseline_image, inspection_image)

        doc.add_heading("4. 종합 의견", level=1)
        doc.add_paragraph(narrative.overall_opinion)

        _add_disclaimer(doc)

        out_dir = output_dir or config.REPORTS_DIR
        out_dir.mkdir(parents=True, exist_ok=True)
        report_path = out_dir / f"inspection_{now:%Y%m%d_%H%M}.docx"
        doc.save(str(report_path))
        logger.info("보고서 생성 완료: %s", report_path)
        return report_path, narrative
